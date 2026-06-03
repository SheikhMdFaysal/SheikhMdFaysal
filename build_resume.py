from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ACCENT = RGBColor(0x1F, 0x4E, 0x79)

doc = Document()
for s in doc.sections:
    s.top_margin = Inches(0.5)
    s.bottom_margin = Inches(0.5)
    s.left_margin = Inches(0.7)
    s.right_margin = Inches(0.7)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)

def add_para(text='', bold=False, size=10.5, color=None, align=None, space_after=2):
    p = doc.add_paragraph()
    if align: p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    return p

def section(title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title.upper())
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = ACCENT
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:color'), '1F4E79')
    pBdr.append(bottom)
    pPr.append(pBdr)

def role_header(left_bold, left_normal, right):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(7.1), WD_ALIGN_PARAGRAPH.RIGHT)
    r = p.add_run(left_bold)
    r.bold = True
    r.font.size = Pt(10.5)
    r2 = p.add_run(left_normal)
    r2.bold = True
    r2.font.size = Pt(10.5)
    r3 = p.add_run('\t' + right)
    r3.bold = True
    r3.font.size = Pt(10.5)

def italic_line(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(10)

def bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.25)
    for run in p.runs:
        run.font.size = Pt(10.5)
    r = p.add_run(text)
    r.font.size = Pt(10.5)

# ---------- HEADER ----------
name = doc.add_paragraph()
name.alignment = WD_ALIGN_PARAGRAPH.CENTER
name.paragraph_format.space_after = Pt(0)
r = name.add_run('SHEIKH MD FAYSAL')
r.bold = True
r.font.size = Pt(20)
r.font.color.rgb = ACCENT

contact = doc.add_paragraph()
contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact.paragraph_format.space_after = Pt(0)
rc = contact.add_run('Jersey City, NJ 07307  •  (201) 526-5280  •  faysals1@montclair.edu  •  linkedin.com/in/faysal-msba  •  github.com/SheikhMdFaysal')
rc.font.size = Pt(10)

# ---------- EDUCATION ----------
section('Education')
role_header('Montclair State University (MSU)', '', 'Jan 2025 – May 2026')
italic_line('Montclair, NJ')
add_para('MS in Business Analytics  |  GPA: 3.83 / 4.0  |  Graduated May 2026', size=10.5)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
r = p.add_run('F-1 OPT Available Jul 2026 — No Sponsorship Required During OPT Period')
r.bold = True
r.font.size = Pt(10.5)
add_para('Relevant Coursework: Data Mining, Data Visualization, Data Wrangling, Database Systems, AI for Business, Optimization Methods, Business Process Management', size=10)

role_header('Shahjalal University of Science & Technology (SUST)', '', '2012 – 2014')
italic_line('Sylhet, Bangladesh')
add_para('MBA, Finance and Banking', size=10.5)

role_header('Shahjalal University of Science & Technology (SUST)', '', '2008 – 2012')
italic_line('Sylhet, Bangladesh')
add_para('BBA, Finance and Banking', size=10.5)

# ---------- SKILLS ----------
section('Technical Skills')
for label, body in [
    ('BI & Visualization: ', 'Power BI (DAX, dashboards), Tableau, Advanced Excel (VLOOKUP, pivot tables, Solver, VBA)'),
    ('Languages & Libraries: ', 'SQL (joins, subqueries, window functions), Python (pandas, NumPy, scikit-learn, matplotlib), R (ggplot2, regression, ANOVA)'),
    ('Analytics Methods: ', 'Exploratory Data Analysis (EDA), Data Wrangling, Preprocessing, Cleaning, Segmentation, Statistics, Regression, ANOVA, Hypothesis Testing, Correlation Analysis, Data Mining, BI, ETL, KPI Development, Forecasting'),
    ('Database & Systems: ', 'RDBMS (PostgreSQL, MySQL, SQL Server), Data Validation, Data Integrity Checks, Process Mining (Celonis), Machine Learning (classification, clustering, neural networks)'),
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    rb = p.add_run(label)
    rb.bold = True
    rb.font.size = Pt(10.5)
    rn = p.add_run(body)
    rn.font.size = Pt(10.5)

# ---------- EXPERIENCE ----------
section('Professional Experience')

role_header('Titas Gas Transmission & Distribution PLC | ', 'Deputy Manager, Financial Planning & Internal Control', 'Dec 2023 – Present')
italic_line('Dhaka, Bangladesh')
for b in [
    'Use SQL and Python to extract, transform, and analyze 3M+ customer records from RDBMS, automating data validation and integrity checks — reducing manual processing by ~200 entries per quarter and improving data quality by 20%.',
    'Build Power BI dashboards (DAX) for financial KPIs, enabling real-time business intelligence reporting across 15+ departments and reducing manual reporting time by 40% (~8 hours/week saved).',
    'Apply regression and time-series forecasting to predict cash flow and budget variances with 85%+ accuracy, supporting data-driven decision making for CFO and VP Finance.',
    'Conduct exploratory data analysis (EDA) and correlation analysis on compliance metrics, achieving 20% improvement in audit accuracy; gather and document business requirements through cross-functional stakeholder collaboration to map financial workflows.',
]: bullet(b)

role_header('Titas Gas Transmission & Distribution PLC | ', 'Assistant Manager, Revenue Collection & Financial Reporting', 'Dec 2018 – Dec 2023')
italic_line('Dhaka, Bangladesh')
for b in [
    'Performed data mining and data segmentation using SQL and Excel on billing patterns, uncovering business trends that drove a 120% increase in revenue recovery ($1.2M+) over 6 months.',
    'Developed interactive Tableau dashboards tracking collection metrics by region and customer segment, delivering business intelligence insights to 3 cross-functional field teams.',
    'Implemented data validation and cleaning frameworks in Python to detect billing errors, improving data quality and reducing disputes by 25% (~$300K saved annually).',
    'Conducted variance analysis identifying $2M+ in uncollected receivables; data-driven recommendations to senior management recovered $800K+ within 90 days.',
]: bullet(b)

role_header('Bank Asia PLC | ', 'Probationary Officer', 'Feb 2016 – Nov 2018')
italic_line('Dhaka, Bangladesh')
for b in [
    'Performed data wrangling and exploratory data analysis using Excel on $10M+ commercial loan portfolio, supporting business analysis of credit risk for 200+ corporate borrowers.',
    'Developed automated credit scoring tools incorporating 15+ financial ratios using Excel (VLOOKUP, pivot tables), reducing loan evaluation time by 30% and improving portfolio data quality (18% reduction in delinquency).',
    'Built Excel dashboards tracking portfolio health metrics (NPL ratio, coverage ratio, vintage analysis), supporting data-driven decision making for senior management.',
]: bullet(b)

role_header('NCC Bank PLC | ', 'Management Trainee Officer', 'Apr 2014 – Feb 2016')
italic_line('Mymensingh, Bangladesh')
for b in [
    'Performed data validation and quality assurance on 500+ monthly KYC records, achieving 99% accuracy and zero compliance violations over 2 years.',
    'Analyzed customer transaction patterns using data segmentation to identify suspicious activities, supporting compliance analytics and AML/CFT regulatory reporting.',
    'Built Excel-based reporting templates improving branch performance reporting efficiency by 25% and standardizing data collection processes across 3 departments.',
]: bullet(b)

# ---------- PROJECTS ----------
section('Projects')

role_header('AI Security Testing Platform', ' | MS Capstone – Ada Analytics | MSU', 'Spring 2026')
bullet('Co-built a web-based platform that safely stress-tests enterprise AI chatbots and large language models (LLMs) before launch — sending hundreds of adversarial prompts to detect data leakage, harmful outputs, and prompt-injection vulnerabilities, then producing a plain-English safety report for business stakeholders.')
bullet('Led business requirements gathering and documentation across 7 Agile sprints; applied data mining and segmentation for market research and built a 3-year financial model supporting a $10M+ Data and Analytics opportunity analysis.')
bullet('Live demo: https://ai-security-platform-jlp76.ondigitalocean.app/  |  Repo: github.com/SheikhMdFaysal/enterprise-ai-security-platform')

role_header('Amazon Beauty Reviews — Text Mining & Sentiment AI', ' | Advanced Data Mining Research (INFO585) – MSU', 'Fall 2025')
bullet('Analyzed 5,200+ Amazon product reviews using TF-IDF document similarity, K-Means and hierarchical clustering, and five supervised ML classifiers (Logistic Regression, Naïve Bayes, SVM, Random Forest, Gradient Boosting).')
bullet('Built and benchmarked CNN and Bidirectional LSTM deep-learning models in TensorFlow/Keras for positive vs. negative sentiment classification on a highly imbalanced dataset (96.2% / 3.8%), turning unstructured review text into actionable business insight.')

role_header('Logistics Workflow Optimization', ' | Process Mining (Celonis) – MSU', 'Fall 2025')
bullet('Used Celonis to perform end-to-end data analysis, identify bottlenecks and process variants, and maintain documentation — reducing simulated cycle time by 22%.')

role_header('Strategic Retail Sales Forecasting', ' | Time-Series Analysis – MSU', 'Spring 2025')
bullet('Used Python and Power BI to conduct EDA, apply data preprocessing and ARIMA/exponential smoothing models for 12-month forecasts — achieving 90%+ accuracy; demonstrated full Business Analyst workflow from data cleaning to interactive dashboard delivery.')

# ---------- VOLUNTEERING & ACHIEVEMENTS ----------
section('Volunteering & Achievements')

role_header('FIFA World Cup 2026™ | ', 'Uniforms Volunteer, NY/NJ Volunteer Center', 'May 2026 – Jul 2026')
italic_line('New York, NY')
for b in [
    'Selected as an official Uniforms Volunteer (Volunteer ID 25412856) supporting the largest sporting event in the world, hosted across the New York / New Jersey region.',
    'Completed venue-specific training and attended multiple operational shifts at the NY/NJ Volunteer Center (1185 6th Ave, NYC), distributing accredited uniforms and kits to thousands of tournament volunteers.',
]: bullet(b)

role_header('Poster Presentation', ' — 2026 Student Research Symposium, Montclair State University', 'May 4, 2026')
bullet('Recognized with a Certificate of Participation issued by the Vice Provost for Research & Interim Head of the Graduate Center for presenting analytics research at the university-wide symposium.')

# ---------- MEMBERSHIPS ----------
section('Professional Memberships')
for m in [
    'Rotary International — Member  |  Rotary ID: 11830255 · Club ID: 225363 · District: 64',
    'INFORMS (Institute for Operations Research and the Management Sciences) — Member  |  Member ID: 1985875 · Paid through Jan 2027',
    'New Jersey Society of Certified Public Accountants (NJCPA) — Student Member  |  ID: 609568 · Hudson Chapter',
]: bullet(m)

# ---------- CERTIFICATIONS ----------
section('Certifications')
for c in [
    'Foundations of Data Ethics – INFORMS  |  Jan 2026',
    'Process Mining: Celonis Academic Fundamentals – Celonis  |  Jun 2025',
    'Process Mining: Celonis Foundations – Celonis  |  Apr 2025',
    'Generative AI Mastermind Workshop – Outskill  |  Nov 2024',
    'Google Crash Course on Python – Coursera/Google  |  Mar 2024',
    'SQL for Data Science – edX/IBM  |  In Progress',
]: bullet(c)

out = r'C:\Users\sober\Documents\Faysal\1st_Project\Resume_SHEIKH MD FAYSAL.docx'
doc.save(out)
print('Saved:', out)
